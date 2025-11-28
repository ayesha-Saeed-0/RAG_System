


import streamlit as st
from langchain_groq import ChatGroq

from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import pandas as pd
import tempfile
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings

import os

# -------------------------------
# STREAMLIT UI
# -------------------------------
st.title("📄 AI Contract Compliance Checker (Groq + LangChain)")
st.write("Upload a contract and automatically check compliance rules.")

# -------------------------------
# ENTER API KEY AT RUNTIME
# -------------------------------
GROQ_API_KEY = st.text_input("Enter your Groq API Key (kept secret)", type="password")

if GROQ_API_KEY:
    # Initialize LLM only after key is provided
    llm = ChatGroq(
        model="llama-3.3-70b-versatile",
        temperature=0,
        groq_api_key=GROQ_API_KEY
    )

    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

    # -------------------------------
    # COMPLIANCE RULES
    # -------------------------------
    COMPLIANCE_RULES = {
        "confidentiality": {
        "rule": "Document must contain confidentiality/NDA clauses",
        "keywords": ["confidential", "proprietary", "trade secret", "non-disclosure", "nda"],
        "severity": "HIGH"
        },
        "termination": {
            "rule": "Clear termination conditions and notice periods must be specified",
            "keywords": ["termination", "notice period", "end of contract", "effective date", "terminate"],
            "severity": "HIGH"
        },
        "liability_limitation": {
            "rule": "Liability limitations must be explicitly defined",
            "keywords": ["liability", "limitation of liability", "exclude", "not liable", "cap on liability"],
            "severity": "HIGH"
        },
        "governing_law": {
            "rule": "Governing law and jurisdiction must be specified",
            "keywords": ["governing law", "jurisdiction", "applicable law", "shall be governed", "laws of"],
            "severity": "HIGH"
        },
        "indemnification": {
            "rule": "Indemnification obligations must be clearly stated",
            "keywords": ["indemnify", "indemnification", "defend", "hold harmless", "harmless from"],
            "severity": "MEDIUM"
        },
        "payment_terms": {
            "rule": "Payment terms, schedule, and methods must be defined",
            "keywords": ["payment", "invoice", "due date", "payment terms", "price", "fee", "billing"],
            "severity": "HIGH"
        },
        "data_protection": {
            "rule": "Data protection and privacy measures must be documented",
            "keywords": ["data protection", "privacy", "gdpr", "personal data", "processing", "data security"],
            "severity": "HIGH"
        },
        "warranty": {
            "rule": "Warranty disclaimers or warranties must be present",
            "keywords": ["warranty", "warrant", "guarantee", "representation", "as-is", "disclaim"],
            "severity": "MEDIUM"
        },
        "dispute_resolution": {
            "rule": "Dispute resolution mechanism must be specified",
            "keywords": ["dispute", "arbitration", "mediation", "litigation", "resolution", "arbitrate"],
            "severity": "MEDIUM"
        },
        "force_majeure": {
            "rule": "Force majeure clause should be included",
            "keywords": ["force majeure", "unforeseen", "circumstances beyond", "act of god", "unforeseeable"],
            "severity": "LOW"
        },
        "renewal_terms": {
            "rule": "Renewal and extension terms must be clear",
            "keywords": ["renew", "extension", "renewal", "continue", "successive", "auto-renew"],
            "severity": "MEDIUM"
        },
        "intellectual_property": {
            "rule": "Intellectual property rights must be addressed",
            "keywords": ["intellectual property", "ip", "copyright", "patent", "ownership", "trademark"],
            "severity": "HIGH"
        },
        "insurance": {
            "rule": "Insurance requirements should be specified",
            "keywords": ["insurance", "insure", "coverage", "premium", "policy", "liability insurance"],
            "severity": "MEDIUM"
        },
        "modification": {
            "rule": "Amendment and modification procedures must be outlined",
            "keywords": ["amendment", "modify", "modification", "change", "waiver", "alter"],
            "severity": "MEDIUM"
        },
        "termination_for_cause": {
            "rule": "Grounds for termination for cause must be defined",
            "keywords": ["termination for cause", "breach", "material breach", "default", "cause"],
            "severity": "HIGH"
        },
        "confidentiality_duration": {
            "rule": "Duration and survival of confidentiality obligations must be specified",
            "keywords": ["confidentiality period", "survival", "years", "duration", "indefinite"],
            "severity": "HIGH"
        },
        "limitation_of_damages": {
            "rule": "Limitation of damages clause should exclude consequential damages",
            "keywords": ["consequential damages", "indirect damages", "lost profits", "damage limitation"],
            "severity": "HIGH"
        },
        "audit_rights": {
            "rule": "Audit and inspection rights must be clearly defined",
            "keywords": ["audit", "inspection", "audit rights", "examine", "verify"],
            "severity": "MEDIUM"
        },
        "insurance_requirements": {
            "rule": "Required types and amounts of insurance must be specified",
            "keywords": ["minimum insurance", "coverage amount", "insurance requirement", "insured"],
            "severity": "MEDIUM"
        },
        "severability": {
            "rule": "Severability clause for invalid provisions should be present",
            "keywords": ["severability", "severability clause", "invalid", "unenforceable"],
            "severity": "MEDIUM"
        },
        "assignment_restriction": {
            "rule": "Restrictions on assignment of rights and obligations must be defined",
            "keywords": ["assignment", "assign", "transfer rights", "not assigned", "consent to assign"],
            "severity": "MEDIUM"
        },
        "notice_requirements": {
            "rule": "Notice provisions and methods of delivery must be specified",
            "keywords": ["notice", "notify", "written notice", "notice to", "delivery of notice"],
            "severity": "MEDIUM"
        },
        "entire_agreement": {
            "rule": "Entire agreement clause should be included",
            "keywords": ["entire agreement", "whole agreement", "supersedes", "prior agreements"],
            "severity": "MEDIUM"
        },
        "counterparts": {
            "rule": "Counterparts and execution provisions should be specified",
            "keywords": ["counterparts", "duplicate originals", "execution", "facsimile", "electronic signature"],
            "severity": "LOW"
        },
        "compliance_laws": {
            "rule": "Compliance with applicable laws and regulations must be required",
            "keywords": ["comply with laws", "applicable laws", "legal requirement", "regulatory"],
            "severity": "HIGH"
        },
        "export_control": {
            "rule": "Export control restrictions should be addressed if applicable",
            "keywords": ["export control", "export restriction", "embargo", "sanctions"],
            "severity": "MEDIUM"
        },
        "subcontracting": {
            "rule": "Subcontracting rights and restrictions must be defined",
            "keywords": ["subcontract", "subcontractor", "third party", "delegate"],
            "severity": "MEDIUM"
        },
        "insurance_coverage": {
            "rule": "Insurance coverage requirements and certificate provision must be specified",
            "keywords": ["certificate of insurance", "insurance certificate", "additional insured"],
            "severity": "MEDIUM"
        },
        "performance_standards": {
            "rule": "Performance standards and service level agreements must be defined",
            "keywords": ["performance standard", "sla", "service level", "performance metric"],
            "severity": "HIGH"
        }
    }

    # -------------------------------
    # FILE UPLOAD
    # -------------------------------
    uploaded_file = st.file_uploader("Upload Contract (PDF, TXT, DOCX)", type=["pdf", "txt", "docx"])

    if uploaded_file:
        st.success("File uploaded!")

        temp_path = os.path.join(tempfile.gettempdir(), uploaded_file.name)
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.read())

        # Extract text
        text = ""
        if uploaded_file.type == "text/plain":
            text = open(temp_path, "r").read()
        elif uploaded_file.type == "application/pdf":
            import PyPDF2
            reader = PyPDF2.PdfReader(temp_path)
            text = "\n".join([page.extract_text() for page in reader.pages])
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            import docx
            doc = docx.Document(temp_path)
            text = "\n".join([p.text for p in doc.paragraphs])

        # Text splitting
        st.info("Splitting text into chunks...")
        splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=200)
        docs = [Document(page_content=t) for t in splitter.split_text(text)]

        # Vector store
        st.info("Creating FAISS embeddings...")
        vector_store = FAISS.from_documents(docs, embedding=embeddings)

        # Compliance check function
        def check_compliance(doc_obj):
            results = []
            for key, rule in COMPLIANCE_RULES.items():
                matched = [kw for kw in rule['keywords'] if kw.lower() in doc_obj.page_content.lower()]
                compliant = "YES" if matched else "NO"

                if compliant == "NO":
                    similar_docs = vector_store.similarity_search(rule["rule"], k=2)
                    context = "\n\n".join([d.page_content for d in similar_docs])
                    prompt = f"""
You are a legal compliance analyst.

Evaluate whether the following contract follows this rule:
Rule: {rule['rule']}

Context (relevant parts of the contract extracted via vector search):
{context}

Respond with:
- Compliance YES/NO
- Missing elements
- Risk level
- Suggest corrections
"""
                    try:
                        llm_response = llm.invoke(prompt).content
                    except Exception as e:
                        llm_response = f"LLM ERROR: {e}"

                    evidence = llm_response
                else:
                    evidence = "Keyword match found. Likely compliant."

                results.append({
                    "Rule": rule['rule'],
                    "Severity": rule['severity'],
                    "Compliant": compliant,
                    "Matched Keywords": ", ".join(matched) if matched else "None",
                    "Evidence": evidence
                })
            return pd.DataFrame(results)

        if st.button("Run Compliance Check"):
            st.warning("Running compliance analysis... This may take 10–20 seconds.")
            df = check_compliance(docs[0])

            st.subheader("📊 Compliance Results")
            st.dataframe(df)

            csv = df.to_csv(index=False).encode("utf-8")
            st.download_button(
                label="⬇ Download Report as CSV",
                data=csv,
                file_name="compliance_report.csv",
                mime="text/csv"
            )
